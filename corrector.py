import streamlit as st
import cv2
import numpy as np
import pandas as pd
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

# ==========================================
# 0. SISTEMA DE LOGIN (SEGURIDAD)
# ==========================================
def check_password():
    """Retorna True si el usuario ingresó la clave correcta"""
    if 'logeado' not in st.session_state:
        st.session_state['logeado'] = False

    if st.session_state['logeado']:
        return True

    # Pantalla de Login
    st.markdown("<h1 style='text-align: center;'>🔐 Acceso al Sistema YACHAY</h1>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1,2,1])
    with col2:
        password = st.text_input("Ingrese la contraseña de administrador:", type="password")
        if st.button("Ingresar al Sistema", type="primary"):
            if password == "456":
                st.session_state['logeado'] = True
                st.rerun()
            else:
                st.error("⛔ Contraseña incorrecta.")
    return False

# ==========================================
# 1. FUNCIÓN: CALCULAR LETRA (AD, A, B, C)
# ==========================================
def obtener_letra_calificacion(nota):
    # Redondeamos para asegurar que 10.5 suba a 11 si fuera necesario, 
    # o usamos la lógica estricta. Aquí usaremos lógica estricta con decimales.
    # Si prefieres redondeo simple: nota = round(nota)
    
    if nota < 0: return "?"
    if nota <= 10.99: return "C"    # 0 a 10 (considerando decimales bajos)
    elif nota <= 13.99: return "B"  # 11 a 13
    elif nota <= 17.99: return "A"  # 14 a 17
    else: return "AD"               # 18 a 20

# ==========================================
# 2. FUNCIÓN: GENERAR HOJA DE RESPUESTAS
# ==========================================
def generar_hoja_examen(num_preguntas):
    width, height = 2480, 3508 
    img = np.ones((height, width, 3), dtype=np.uint8) * 255
    COLOR_NEGRO = (0, 0, 0)
    
    # Encabezado Hoja
    cv2.putText(img, "IEP YACHAY - HOJA DE RESPUESTAS", (600, 200), cv2.FONT_HERSHEY_SIMPLEX, 2.5, COLOR_NEGRO, 5)
    cv2.putText(img, "EDUCAR PARA LA VIDA", (950, 300), cv2.FONT_HERSHEY_SIMPLEX, 1.5, COLOR_NEGRO, 3)
    cv2.putText(img, "Nombre: ___________________________________________________", (200, 500), cv2.FONT_HERSHEY_SIMPLEX, 1.5, COLOR_NEGRO, 3)
    
    start_y, start_x = 900, 300
    spacing_y, column_spacing = 110, 700
    preguntas_por_columna = 20
    
    for i in range(num_preguntas):
        columna_actual = i // preguntas_por_columna
        fila_en_columna = i % preguntas_por_columna
        x_base = start_x + (columna_actual * column_spacing)
        y_base = start_y + (fila_en_columna * spacing_y)
        
        cv2.putText(img, f"{i+1}.", (x_base - 120, y_base + 25), cv2.FONT_HERSHEY_SIMPLEX, 1.3, COLOR_NEGRO, 3)
        letras = ['A', 'B', 'C', 'D']
        for j, letra in enumerate(letras):
            cx = x_base + (j * 130)
            cy = y_base - 10
            cv2.circle(img, (cx, cy), 40, COLOR_NEGRO, 7)
            cv2.putText(img, letra, (cx - 18, cy + 18), cv2.FONT_HERSHEY_SIMPLEX, 1.2, COLOR_NEGRO, 3)
            
    return img

# ==========================================
# 3. FUNCIÓN: PROCESAR IMAGEN
# ==========================================
def procesar_imagen(img_bytes, total_preguntas_esperadas):
    file_bytes = np.asarray(bytearray(img_bytes), dtype=np.uint8)
    img = cv2.imdecode(file_bytes, 1)
    gris = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)
    blur = cv2.GaussianBlur(gris, (5, 5), 0)
    bordes = cv2.Canny(blur, 75, 200)
    cnts, _ = cv2.findContours(bordes, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
    
    burbujas = []
    for c in cnts:
        (x, y, w, h) = cv2.boundingRect(c)
        if w >= 20 and h >= 20 and 0.8 <= w/float(h) <= 1.2:
            burbujas.append(c)
    
    burbujas = sorted(burbujas, key=lambda c: cv2.boundingRect(c)[1])
    
    filas = []
    if burbujas:
        fila_act = [burbujas[0]]
        for i in range(1, len(burbujas)):
            if abs(cv2.boundingRect(burbujas[i])[1] - cv2.boundingRect(fila_act[-1])[1]) < 30:
                fila_act.append(burbujas[i])
            else:
                filas.append(sorted(fila_act, key=lambda c: cv2.boundingRect(c)[0]))
                fila_act = [burbujas[i]]
        filas.append(sorted(fila_act, key=lambda c: cv2.boundingRect(c)[0]))

    respuestas_alumno = []
    _, bin_inv = cv2.threshold(gris, 140, 255, cv2.THRESH_BINARY_INV)
    
    for q, fila in enumerate(filas):
        if q >= total_preguntas_esperadas: break 
        if len(fila) >= 1:
            max_p = 0
            idx_marcado = -1
            for k, c in enumerate(fila):
                mask = np.zeros(gris.shape, dtype="uint8")
                cv2.drawContours(mask, [c], -1, 255, -1)
                count = cv2.countNonZero(cv2.bitwise_and(bin_inv, bin_inv, mask=mask))
                if count > max_p:
                    max_p = count
                    idx_marcado = k
            
            opciones = ['A', 'B', 'C', 'D']
            if max_p > 100 and idx_marcado < 4:
                respuestas_alumno.append(opciones[idx_marcado])
            else:
                respuestas_alumno.append('X') 
    
    return respuestas_alumno

# ==========================================
# 4. FUNCIÓN: CREAR WORD OFICIAL IEP YACHAY
# ==========================================
def crear_reporte_word(df, grado, nivel):
    doc = Document()
    fecha_hoy = datetime.now().strftime("%d/%m/%Y")

    # ENCABEZADO
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INSTITUCIÓN EDUCATIVA PARTICULAR YACHAY\n")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0, 51, 102) 
    
    run2 = p.add_run("PIONEROS EN LA EDUCACIÓN DE CALIDAD")
    run2.italic = True
    run2.font.size = Pt(11)
    
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_info = p_info.add_run(f"NIVEL: {nivel.upper()}  |  GRADO: {grado.upper()}  |  FECHA: {fecha_hoy}")
    run_info.bold = True
    
    doc.add_paragraph('--------------------------------------------------').alignment = WD_ALIGN_PARAGRAPH.CENTER

    t_titulo = doc.add_heading('RANKING DE MÉRITOS', level=1)
    t_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # TABLA
    cols_fijas = ['PUESTO', 'ALUMNO']
    # Filtramos columnas que no son cursos
    cols_cursos = [c for c in df.columns if c not in ['Puesto', 'Alumno', 'Promedio Final']]
    cols_finales = cols_fijas + cols_cursos + ['PROMEDIO FINAL']
    
    t = doc.add_table(rows=1, cols=len(cols_finales))
    t.style = 'Table Grid'
    
    # Encabezados
    hdr_cells = t.rows[0].cells
    for i, col_name in enumerate(cols_finales):
        hdr_cells[i].text = col_name
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # Llenar datos
    for index, row in df.iterrows():
        row_cells = t.add_row().cells
        row_cells[0].text = str(row['Puesto'])
        row_cells[1].text = str(row['Alumno'])
        
        for i, curso in enumerate(cols_cursos):
            row_cells[2 + i].text = str(row[curso])
            
        row_cells[-1].text = str(row['Promedio Final'])
    
    doc.add_paragraph('\n')
    
    # PIE DE PÁGINA
    p_final = doc.add_paragraph()
    p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_final = p_final.add_run("“EDUCAR PARA LA VIDA”")
    run_final.bold = True
    run_final.font.size = Pt(14)
    
    buffer_word = io.BytesIO()
    doc.save(buffer_word)
    buffer_word.seek(0)
    return buffer_word

# ==========================================
# 5. INTERFAZ GRÁFICA PRINCIPAL
# ==========================================
st.set_page_config(page_title="Intranet IEP YACHAY", page_icon="🏫", layout="wide")

# CHECK LOGIN
if not check_password():
    st.stop()

if 'cursos' not in st.session_state:
    st.session_state['cursos'] = []

st.markdown("<h1 style='text-align: center; color: #003366;'>IEP YACHAY - EDUCAR PARA LA VIDA 🏫</h1>", unsafe_allow_html=True)
st.markdown("<h3 style='text-align: center;'>Sistema Integrado de Calificación y Ranking</h3>", unsafe_allow_html=True)
st.divider()

with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/2997/2997259.png", width=100)
    st.header("⚙️ Configuración Global")
    
    modo = st.radio("Seleccione Módulo:", ["1. Corregir Examen", "2. Generar Hoja Óptica"])
    
    st.divider()
    st.subheader("🎓 Datos Académicos")
    nivel = st.selectbox("Nivel Educativo:", ["Inicial", "Primaria", "Secundaria", "Pre-Universitario"])
    
    grados_opt = []
    if nivel == "Inicial": grados_opt = ["3 años", "4 años", "5 años"]
    elif nivel == "Primaria": grados_opt = [f"{i}° Grado" for i in range(1, 7)]
    elif nivel == "Secundaria": grados_opt = [f"{i}° Año" for i in range(1, 6)]
    else: grados_opt = ["Pre-U A", "Pre-U B", "Intensivo"]
    
    grado = st.selectbox("Grado / Salón:", grados_opt)
    
    st.divider()
    if st.button("🔒 Cerrar Sesión"):
        st.session_state['logeado'] = False
        st.rerun()

# ==========================================
# MÓDULO 2: GENERAR HOJA
# ==========================================
if modo == "2. Generar Hoja Óptica":
    st.header("🖨️ Centro de Impresión")
    n_preg = st.slider("Cantidad total de preguntas en la hoja:", 10, 100, 20)
    if st.button("Generar Hoja Personalizada PDF/JPG"):
        img = generar_hoja_examen(n_preg)
        is_success, buffer = cv2.imencode(".jpg", img)
        st.download_button("⬇️ Descargar Hoja para Imprimir", io.BytesIO(buffer), "Hoja_YACHAY.jpg", "image/jpeg")

# ==========================================
# MÓDULO 1: CORRECCIÓN MULTI-CURSO
# ==========================================
elif modo == "1. Corregir Examen":
    
    col_izq, col_der = st.columns([1, 2])
    
    with col_izq:
        st.subheader("1. Configurar Examen")
        
        # Input del total de preguntas de la hoja física
        total_hoja = st.number_input("Total preguntas en la Hoja Física:", min_value=5, max_value=100, value=20)
        
        # CÁLCULOS
        usadas = sum([len(c['claves']) for c in st.session_state['cursos']])
        disponibles = total_hoja - usadas
        
        st.markdown("### 📊 Estado de la Hoja")
        col_a, col_b = st.columns(2)
        col_a.metric("Preguntas Usadas", usadas)
        col_b.metric("Preguntas Disponibles", disponibles)

        st.write("---")
        st.markdown("#### ➕ Agregar Nuevo Curso")
        
        if disponibles > 0:
            nuevo_inicio = usadas + 1
            st.info(f"💡 El próximo curso comenzará en la **Pregunta N° {nuevo_inicio}**")
            
            nuevo_curso = st.text_input("Nombre del Curso (ej: Matemática)")
            nuevas_claves = st.text_area("Claves (ej: AAAB...)", height=80, placeholder="Escribe las claves aquí...").upper().strip().replace(" ", "").replace("\n", "")
            
            # Cálculo anticipado
            if nuevas_claves:
                cant_nuevas = len(nuevas_claves)
                fin_rango = nuevo_inicio + cant_nuevas - 1
                if cant_nuevas <= disponibles:
                    st.success(f"✅ Este curso cubrirá de la **{nuevo_inicio}** a la **{fin_rango}**.")
                else:
                    st.error(f"❌ ¡Te pasaste! Intentas poner {cant_nuevas} preguntas pero solo quedan {disponibles}.")

            if st.button("Agregar Área"):
                if not nuevo_curso:
                    st.error("⚠️ Falta el nombre del curso.")
                elif len(nuevas_claves) == 0:
                    st.error("⚠️ Faltan las claves.")
                elif len(nuevas_claves) > disponibles:
                    st.error(f"⚠️ Error: Solo quedan {disponibles} espacios libres.")
                else:
                    st.session_state['cursos'].append({'nombre': nuevo_curso, 'claves': nuevas_claves})
                    st.toast(f"✅ Curso {nuevo_curso} agregado correctamente.")
                    st.rerun()
        else:
            st.warning("🔒 La hoja está llena. No puedes agregar más cursos.")
            if st.button("🗑️ Borrar todo y empezar de cero"):
                st.session_state['cursos'] = []
                st.rerun()

        # TABLA RESUMEN
        if st.session_state['cursos']:
            st.write("---")
            st.caption("Estructura actual del examen:")
            resumen = []
            inicio_temp = 1
            for c in st.session_state['cursos']:
                fin_temp = inicio_temp + len(c['claves']) - 1
                resumen.append({
                    "Curso": c['nombre'],
                    "Preguntas": len(c['claves']),
                    "Rango": f"{inicio_temp} al {fin_temp}"
                })
                inicio_temp = fin_temp + 1
            
            st.table(pd.DataFrame(resumen))

    with col_der:
        st.subheader("2. Procesamiento de Exámenes")
        
        if not st.session_state['cursos']:
            st.info("👈 Primero configura los cursos en el panel izquierdo.")
        else:
            files = st.file_uploader("Sube las fotos de los alumnos:", accept_multiple_files=True)
            
            if files and st.button("🏆 CALIFICAR Y GENERAR RANKING", type="primary"):
                
                datos_numericos = []
                total_preguntas_examen = sum([len(c['claves']) for c in st.session_state['cursos']])
                
                barra = st.progress(0, text="Procesando imágenes...")
                
                for idx, file in enumerate(files):
                    try:
                        nombre_alumno = file.name.split('.')[0]
                        file_bytes = file.read()
                        respuestas_detectadas = procesar_imagen(file_bytes, total_preguntas_examen)
                        
                        idx_inicio = 0
                        alumno_data = {"Alumno": nombre_alumno}
                        suma_notas = 0
                        
                        # CALCULO NUMÉRICO PURO (Base 20)
                        for curso in st.session_state['cursos']:
                            nombre_area = curso['nombre']
                            claves_area = curso['claves']
                            cant_preg = len(claves_area)
                            
                            respuestas_curso = respuestas_detectadas[idx_inicio : idx_inicio + cant_preg]
                            aciertos = 0
                            for r_alum, r_clave in zip(respuestas_curso, claves_area):
                                if r_alum == r_clave:
                                    aciertos += 1
                            
                            nota_curso = round((aciertos / cant_preg) * 20, 2) if cant_preg > 0 else 0
                            alumno_data[nombre_area] = nota_curso
                            suma_notas += nota_curso
                            idx_inicio += cant_preg
                        
                        promedio_final = round(suma_notas / len(st.session_state['cursos']), 2)
                        alumno_data["Promedio Final"] = promedio_final
                        
                        datos_numericos.append(alumno_data)
                        
                    except Exception as e:
                        st.error(f"Error en {file.name}: {e}")
                    
                    barra.progress((idx + 1) / len(files))
                
                barra.empty()
                
                if datos_numericos:
                    # 1. Crear DataFrame Numérico para ordenar
                    df_num = pd.DataFrame(datos_numericos)
                    df_num = df_num.sort_values(by="Promedio Final", ascending=False).reset_index(drop=True)
                    df_num.insert(0, 'Puesto', df_num.index + 1)
                    
                    # 2. Crear DataFrame de Visualización (Con letras si aplica)
                    df_vis = df_num.copy()
                    
                    # APLICAR LÓGICA DE LETRAS SOLO SI ES INICIAL, PRIMARIA O SECUNDARIA
                    usar_letras = nivel in ["Inicial", "Primaria", "Secundaria"]
                    
                    if usar_letras:
                        cols_a_convertir = [c for c in df_vis.columns if c not in ['Puesto', 'Alumno']]
                        for col in cols_a_convertir:
                            # Convertimos cada nota numérica al formato "Nota (Letra)"
                            df_vis[col] = df_vis[col].apply(lambda x: f"{x} ({obtener_letra_calificacion(x)})")

                    st.balloons()
                    st.success(f"✅ ¡Proceso Terminado! Se calificaron {len(files)} alumnos.")
                    
                    st.subheader(f"🏆 Ranking Oficial: {nivel} - {grado}")
                    
                    # Mostrar tabla
                    st.dataframe(df_vis)
                    
                    # Generar Word con el DF visual (que ya tiene las letras)
                    word_bytes = crear_reporte_word(df_vis, grado, nivel)
                    
                    st.download_button(
                        label="📄 DESCARGAR REPORTE OFICIAL (WORD)",
                        data=word_bytes,
                        file_name=f"Ranking_YACHAY_{grado}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )